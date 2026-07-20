import streamlit as st
import requests
import PyPDF2
from io import BytesIO
import base64
from docx import Document
from fpdf import FPDF
import pandas as pd

st.set_page_config(layout="wide", page_title="Agentic Proposal Evaluator")

API_URL = "http://localhost:8000"

def extract_text_from_pdf(file):
    try:
        reader = PyPDF2.PdfReader(BytesIO(file.read()))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error extracting PDF: {e}"

def extract_excel_base64(file):
    try:
        bytes_data = file.read()
        return base64.b64encode(bytes_data).decode('utf-8')
    except Exception as e:
        return ""

st.title("🤖 Agentic Business Proposal Evaluator")

with st.sidebar:
    st.header("Navigation")
    if st.button("➕ New Submission"):
        st.session_state["view"] = "new"
    if st.button("⚙️ Admin - Criteria Setup"):
        st.session_state["view"] = "admin"
        
    st.markdown("---")
    st.header("Submission History")
    
    try:
        res = requests.get(f"{API_URL}/history")
        if res.status_code == 200:
            history = res.json()
            if not history:
                st.write("No history yet.")
            for item in history:
                try:
                    title = item.get("payload", {}).get("evaluation_result", {}).get("structured_data", {}).get("summary", "Proposal")[:30] + "..."
                except:
                    title = f"Report {item['id'][:8]}"
                    
                if st.button(f"📄 {title}", key=item["id"]):
                    st.session_state["view"] = "report"
                    st.session_state["current_report"] = item["payload"]["evaluation_result"]
    except Exception as e:
        st.error(f"Could not connect to backend: {e}")

if "view" not in st.session_state:
    st.session_state["view"] = "new"

if st.session_state["view"] == "admin":
    st.header("Admin - Criteria Setup")
    st.markdown("Upload a TXT/PDF file or type the criteria below.")
    
    criteria_file = st.file_uploader("Upload Criteria File (PDF/TXT)", type=["pdf", "txt"])
    
    # Try to load existing criteria
    existing_criteria = ""
    try:
        res = requests.get(f"{API_URL}/criteria")
        if res.status_code == 200:
            existing_criteria = res.json().get("criteria_text", "")
    except:
        pass
        
    criteria_text = st.text_area("Criteria Details", value=existing_criteria, height=200)
    
    if st.button("Save Criteria", type="primary"):
        if criteria_file:
            with st.spinner("Uploading and processing criteria file..."):
                files_to_upload = [("files", (criteria_file.name, criteria_file.getvalue(), criteria_file.type))]
                try:
                    upload_resp = requests.post(f"{API_URL}/upload_criteria", files=files_to_upload)
                    if upload_resp.status_code == 200:
                        st.success("Criteria file saved to Knowledge Base successfully!")
                    else:
                        st.error(f"Failed to save criteria file: {upload_resp.text}")
                except Exception as e:
                    st.error(f"Failed to connect to backend: {e}")
        else:
            try:
                resp = requests.post(f"{API_URL}/criteria", json={"criteria_text": criteria_text})
                if resp.status_code == 200:
                    st.success("Criteria text saved successfully!")
                else:
                    st.error("Failed to save criteria text.")
            except Exception as e:
                st.error(f"Failed to connect to backend: {e}")

elif st.session_state["view"] == "new":
    col_left, col_right = st.columns(2)
    
    with col_left:
        st.header("Submit New Proposal")
        uploaded_files = st.file_uploader("Users can Upload Files in PDF/Excel format", type=["pdf", "xlsx", "xls"], accept_multiple_files=True)
        upload_btn = st.button("Upload to Knowledge Base", key="main_upload_btn", type="primary")
        
    with col_right:
        st.header("Upload Previous Proposals")
        knowledge_files = st.file_uploader("Upload Previous Proposals (PDF, Excel)", type=["pdf", "xlsx", "xls"], accept_multiple_files=True, key="knowledge_files")
        kb_upload_btn = st.button("Upload to Knowledge Base", key="kb_upload_btn", type="primary")

    st.markdown("---")
    eval_disabled = not st.session_state.get("upload_success", False)
    evaluate_btn = st.button("Evaluate Proposal", type="primary", disabled=eval_disabled)

    if upload_btn:
        if not uploaded_files:
            st.warning("Please upload at least one file.")
        else:
            with st.spinner("Uploading files..."):
                files_to_upload = [("files", (f.name, f.getvalue(), f.type)) for f in uploaded_files]
                try:
                    upload_resp = requests.post(f"{API_URL}/upload_knowledge", files=files_to_upload)
                    if upload_resp.status_code == 200:
                        st.session_state["proposal_id"] = upload_resp.json().get("proposal_id", "default")
                        st.session_state["upload_success"] = True
                        st.success("Files uploaded successfully! You can now Evaluate the Proposal.")
                        st.rerun()
                    else:
                        st.error(f"Error uploading files: {upload_resp.text}")
                except Exception as e:
                    st.error(f"Failed to connect to backend: {e}")

    if evaluate_btn:
        with st.spinner("Processing evaluation..."):
            proposal_id = st.session_state.get("proposal_id")
            payload = {"proposal_id": proposal_id}
            try:
                resp = requests.post(f"{API_URL}/evaluate", json=payload)
                if resp.status_code == 200:
                    data = resp.json()
                    st.session_state["current_report"] = data["result"]
                    st.session_state["show_report_in_new"] = True
                else:
                    st.error(f"Error from backend evaluate: {resp.text}")
            except Exception as e:
                st.error(f"Failed to connect to backend: {e}")
        
    if kb_upload_btn:
        if not knowledge_files:
            st.warning("Please upload at least one file.")
        else:
            with st.spinner("Uploading and processing knowledge base files..."):
                files_to_upload = [("files", (f.name, f.getvalue(), f.type)) for f in knowledge_files]
                try:
                    upload_resp = requests.post(f"{API_URL}/upload_knowledge", files=files_to_upload)
                    if upload_resp.status_code == 200:
                        st.success("Successfully uploaded to Knowledge Base!")
                    else:
                        st.error(f"Error uploading files: {upload_resp.text}")
                except Exception as e:
                    st.error(f"Failed to connect to backend: {e}")

    if st.session_state.get("show_report_in_new"):
        st.markdown("---")
        st.header("Evaluation Report")
        report_data = st.session_state.get("current_report", {})
        
        tab1, tab2, tab3, tab4 = st.tabs([
            "Criteria Evaluation", 
            "Division", 
            "Evaluation score", 
            "Executive Report"
        ])
        
        with tab1:
            st.subheader("Criteria Evaluation Results")
            criteria_scores = report_data.get("criteria_scores", {})
            if isinstance(criteria_scores, dict) and "criteria_evaluations" in criteria_scores:
                evals = criteria_scores["criteria_evaluations"]
                df = pd.DataFrame(evals)
                if not df.empty:
                    df = df.rename(columns={"criteria": "Criteria", "status": "Result", "reason": "Assessment"})
                    if "Implication" not in df.columns:
                        df["Implication"] = "N/A"
                    cols = [c for c in ["Criteria", "Result", "Assessment", "Implication"] if c in df.columns]
                    st.table(df[cols])
            else:
                st.write(criteria_scores)

        with tab2:
            st.subheader("Division Classification")
            print("Division", report_data.get("classification", {}))
            st.json(report_data.get("classification", {}))
            
        with tab3:
            st.subheader("Evaluation Score")
            st.markdown(report_data.get("final_score", "No report generated."))
            with st.expander("View Raw Agent Outputs"):
                st.json(report_data)
                
        with tab4:
            st.subheader("Summary Report")
            st.markdown(report_data.get("final_report", "No report generated."))
            def create_docx(text):
                doc = Document()
                doc.add_heading('Executive Summary', 0)
                doc.add_paragraph(text)
                bio = BytesIO()
                doc.save(bio)
                return bio.getvalue()

            print("executive report ", report_data)    
            summary_text = report_data.get("structured_data", {}).get("summary", "No executive summary available.")
            try:
                docx_data = create_docx(summary_text)
                st.download_button(
                    label="Download Executive Summary (DOCX)",
                    data=docx_data,
                    file_name="summary_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Failed to create DOCX: {e}")


elif st.session_state["view"] == "report":
    st.header("Evaluation Report")
    
    report_data = st.session_state.get("current_report", {})
    
    score = report_data.get("final_score", 0)
    rating = report_data.get("rating", "N/A")
    st.metric(label="Opportunity Score", value=f"{score:.2f} / 10", delta=rating)
    
    st.markdown("---")
    
    final_report_md = report_data.get("final_report", "No report generated.")
    st.markdown(final_report_md)
    
    with st.expander("View Raw Agent Outputs"):
        st.json(report_data)
